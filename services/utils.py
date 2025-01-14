import os
from django.core.exceptions import ValidationError
from docx import Document
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import subprocess
from authentication.utils import create_customer

import io
import base64
import gspread
from collections import Counter
import matplotlib

matplotlib.use('Agg')
import matplotlib.pyplot as plt
from django.utils.html import format_html
from exceptions.error_messages import ErrorCodes
from exceptions.exception import CustomApiException
from google.oauth2.service_account import Credentials


def validate_file_type_and_size(value):
    """
    Validates the file type and size of uploaded files.

    Allowed types: PDF, Word, Excel
    Maximum size: 5MB
    """
    allowed_extensions = ['.pdf', '.mp4', '.mov', '.avi', '.mkv', '.flv', '.ppt', '.pptx']
    max_file_size = 50 * 1024 * 1024  # 50 MB

    # Get file extension
    ext = os.path.splitext(value.name)[1].lower()

    if ext not in allowed_extensions:
        raise ValidationError(f"Unsupported file extension. Allowed types are: {', '.join(allowed_extensions)}")

    # Check file size
    if value.size > max_file_size:
        raise ValidationError(f"The file size exceeds the 50MB limit.")


def center_cell_text(cell, left_indent=4.0, right_indent=4.0):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.paragraph_format.left_indent = Pt(left_indent)
        paragraph.paragraph_format.right_indent = Pt(right_indent)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def center_left_paragraph(paragraph, left_indent=4.0, right_indent=4.0):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.left_indent = Pt(left_indent)
    paragraph.paragraph_format.right_indent = Pt(right_indent)


def file_one_create(serialized_data):
    file_path = "/home/otabek/Desktop/1-variant uchun shakl.docx"
    document = Document(file_path)

    paragraph1 = document.paragraphs[0]
    paragraph2 = document.paragraphs[1]

    paragraph1.text = f'{serialized_data.get('organization_director_position')} {serialized_data.get('organization_director_full_name')}гa'
    paragraph2.text = f'{serialized_data.get('employee_position')} {serialized_data.get('employee_full_name')}дан\n'

    run1 = paragraph1.runs[0]
    run2 = paragraph2.runs[0]

    run1.font.size = Pt(14)
    run2.font.size = Pt(14)

    center_left_paragraph(paragraph=paragraph1, left_indent=300, right_indent=2)
    center_left_paragraph(paragraph=paragraph2, left_indent=300, right_indent=2)

    table1 = document.tables[0]
    table2 = document.tables[1]
    table3 = document.tables[2]
    table4 = document.tables[3]
    table5 = document.tables[4]

    table1.cell(0, 2).text = (serialized_data.get('employee_passport_series') or ''
                              + '\n' + serialized_data.get('employee_passport_taken_date'))
    table1.cell(1, 2).text = serialized_data.get('employee_passport_number') or ''
    table2.cell(1, 2).text = serialized_data.get('related_persons_full_name') or ''
    table2.cell(2, 2).text = (serialized_data.get('related_persons_passport_series') or ''
                              + '\n' + serialized_data.get('related_persons_passport_taken_date') or '')
    table2.cell(5, 2).text = serialized_data.get('employee_legal_entity_name') or ''
    table2.cell(3, 2).text = serialized_data.get('related_persons_passport_number') or ''
    table2.cell(6, 2).text = serialized_data.get('employee_stir_number') or ''
    table2.cell(8, 2).text = serialized_data.get('related_persons_legal_entity_name') or ''
    table2.cell(9, 2).text = serialized_data.get('related_persons_stir_number') or ''
    table3.cell(0, 2).text = serialized_data.get('description') or ''
    table4.cell(0, 0).text = serialized_data.get('employee_position') or ''
    table4.cell(0, 1).text = f"________________________\n Шахсий имзо ёки электрон рақамли имзоси"
    table4.cell(0,
                2).text = f"{serialized_data.get('employee_full_name') or ''} \n {serialized_data.get('filled_date') or ''}"
    table5.cell(0, 0).text = serialized_data.get('organization_director_position') or ''
    table5.cell(0, 1).text = f"\n_______________________\n Шахсий имзо ёки электрон рақамли имзоси"
    table5.cell(0, 2).text = (f"{serialized_data.get('organization_director_full_name') or ''}"
                              f" \n \n _____________________\nТўлдирилган сана")

    center_cell_text(table1.cell(0, 2))
    center_cell_text(table1.cell(1, 2))
    center_cell_text(table2.cell(1, 2))
    center_cell_text(table2.cell(2, 2))
    center_cell_text(table2.cell(3, 2))
    center_cell_text(table2.cell(5, 2))
    center_cell_text(table2.cell(6, 2))
    center_cell_text(table2.cell(8, 2))
    center_cell_text(table2.cell(9, 2))
    center_cell_text(table3.cell(0, 2))
    center_cell_text(table4.cell(0, 0))
    center_cell_text(table4.cell(0, 1))
    center_cell_text(table4.cell(0, 2))
    center_cell_text(table5.cell(0, 0))
    center_cell_text(table5.cell(0, 1))
    center_cell_text(table5.cell(0, 2))

    docx_output_path = f"/home/otabek/Desktop/file1_{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}.docx"
    document.save(docx_output_path)
    pdf_output_path = docx_output_path.replace('.docx', '.pdf')
    output_dir = '/'.join(pdf_output_path.split('/')[:-1])
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_output_path],
                   check=True)
    return serialized_data


def file_two_create(serialized_data):
    file_path = '/home/otabek/Desktop/2-variant uchun shakl.docx'
    document = Document(file_path)

    paragraph0 = document.paragraphs[0]
    paragraph1 = document.paragraphs[1]
    paragraph2 = document.paragraphs[2]

    paragraph0.text = "Ходимнинг эҳтимолий манфаатлар тўқнашуви тўғрисидаги"
    run0 = paragraph0.runs[0]
    run0.font.size = Pt(14)
    paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph0.runs[0].bold = True

    paragraph1.text = "ДЕКЛАРАЦИЯ"
    run1 = paragraph1.runs[0]
    run1.font.size = Pt(14)
    paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph1.runs[0].bold = True

    paragraph2.text = f"Мен, {serialized_data.get('employee_full_name')} ушбу тўлдираётган декларацияда ўзим ва менга алоқадор шахсларнинг эҳтимолий манфаатлар тўқнашувига оид қуйидаги маълумотларни ошкор қиламан:"
    run2 = paragraph2.runs[0]
    run2.font.size = Pt(12)

    table1 = document.tables[0]
    table2 = document.tables[1]
    table3 = document.tables[2]

    table1.cell(0, 2).text = (serialized_data.get('employee_passport_series') or '' +
                              '\n' + serialized_data.get('employee_passport_taken_date') or '')
    table1.cell(1, 2).text = serialized_data.get('employee_passport_number') or ''
    table2.cell(1, 2).text = serialized_data.get('related_persons_full_name') or ''
    table2.cell(2, 2).text = (serialized_data.get('related_persons_passport_series') or '' +
                              '\n' + serialized_data.get('related_persons_passport_taken_date') or '')
    table2.cell(3, 2).text = serialized_data.get('related_persons_passport_number') or ''
    table2.cell(5, 2).text = serialized_data.get('employee_legal_entity_name') or ''
    table2.cell(6, 2).text = serialized_data.get('employee_stir_number') or ''
    table2.cell(8, 2).text = serialized_data.get('related_persons_legal_entity_name') or ''
    table2.cell(9, 2).text = serialized_data.get('related_persons_stir_number') or ''
    table3.cell(0, 0).text = serialized_data.get('employee_position') or ''
    table3.cell(0, 1).text = f"________________________\n Шахсий имзо ёки электрон рақамли имзоси"
    table3.cell(0,
                2).text = f"{serialized_data.get('employee_full_name') or ''} \n {serialized_data.get('filled_date') or ''}"

    center_cell_text(table1.cell(0, 2))
    center_cell_text(table1.cell(1, 2))
    center_cell_text(table2.cell(1, 2))
    center_cell_text(table2.cell(2, 2))
    center_cell_text(table2.cell(3, 2))
    center_cell_text(table2.cell(5, 2))
    center_cell_text(table2.cell(6, 2))
    center_cell_text(table2.cell(8, 2))
    center_cell_text(table2.cell(9, 2))
    center_cell_text(table3.cell(0, 0))
    center_cell_text(table3.cell(0, 1))
    center_cell_text(table3.cell(0, 2))

    docx_output_path = f"/home/otabek/Desktop/file2_{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}.docx"
    document.save(docx_output_path)
    pdf_output_path = docx_output_path.replace('.docx', '.pdf')
    output_dir = '/'.join(pdf_output_path.split('/')[:-1])
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_output_path],
                   check=True)
    return serialized_data


def file_three_create(serialized_data):
    file_path = '/home/otabek/Desktop/3-variant uchun shakl.docx'
    document = Document(file_path)

    paragraph0 = document.paragraphs[0]
    paragraph1 = document.paragraphs[1]
    paragraph2 = document.paragraphs[2]

    paragraph0.text = "Алоқадор шахсларнинг эҳтимолий манфаатлар тўқнашуви тўғрисидаги"
    paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph0.runs[0].bold = True

    paragraph1.text = "ДЕКЛАРАЦИЯ"
    paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph1.runs[0].bold = True

    paragraph2.text = f"Мен {serialized_data.get("related_persons_full_name") or ''}, ушбу декларацияда алоқадор шахс сифатида ўзим ва ходим (ишга кираётган номзоднинг) эҳтимолий манфаатлар тўқнашувига оид қуйидаги маълумотларни маълум қиламан:\n"
    paragraph2.runs[0].bold = False

    table1 = document.tables[0]
    table2 = document.tables[1]

    table1.cell(0, 2).text = serialized_data.get('related_persons_passport_number') or ''
    table1.cell(1, 2).text = serialized_data.get('employee_legal_entity_name') or ''
    table1.cell(2, 2).text = serialized_data.get('employee_stir_number') or ''
    table1.cell(3, 2).text = (serialized_data.get('employee_full_name') or ''
                              + serialized_data.get('employee_position') or '')
    table1.cell(4, 2).text = serialized_data.get('related_persons_kinship_data') or ''
    table1.cell(5, 2).text = serialized_data.get('employee_legal_entity_data') or ''

    table2.cell(0, 1).text = f"________________________\n Шахсий имзо ёки электрон рақамли имзоси"
    table2.cell(0, 2).text = (f"{serialized_data.get('related_persons_full_name') or ''}"
                              f" \n {serialized_data.get('filled_date') or ''}")

    center_cell_text(table1.cell(0, 2))
    center_cell_text(table1.cell(1, 2))
    center_cell_text(table1.cell(2, 2))
    center_cell_text(table1.cell(3, 2))
    center_cell_text(table1.cell(4, 2))
    center_cell_text(table1.cell(5, 2))
    center_cell_text(table2.cell(0, 1))
    center_cell_text(table2.cell(0, 2))

    docx_output_path = f"/home/otabek/Desktop/file3_{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}.docx"
    document.save(docx_output_path)
    pdf_output_path = docx_output_path.replace('.docx', '.pdf')
    output_dir = '/'.join(pdf_output_path.split('/')[:-1])
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_output_path],
                   check=True)
    return serialized_data


def calculate_percent(customer, category_id: int):
    from .models import HonestyTest, HonestyTestResult

    all_tests = HonestyTest.objects.filter(category_id=category_id).count()
    true_tests = HonestyTestResult.objects.filter(test__category_id=category_id, customer_id=customer.id, result=True).count()
    return (true_tests / all_tests) * 100 or 0


def get_google_sheet_statistics(sheets_id):
    base_dir = os.path.dirname(__file__)
    os.makedirs(os.path.join(base_dir, 'google_access_json'), exist_ok=True)
    if not os.listdir(f'{base_dir}/google_access_json'):
        raise CustomApiException(ErrorCodes.INVALID_INPUT, message='Google Sheet access json folder not found.')

    SCOPE = ["https://www.googleapis.com/auth/spreadsheets"]
    creds = Credentials.from_service_account_file(
        filename=f"{base_dir}/google_access_json/oceanic-catcher-441013-i5-32aa559cd203.json", scopes=SCOPE)
    client = gspread.authorize(creds)
    rows = client.open_by_key(sheets_id).sheet1.get_all_values()

    if not rows:
        raise CustomApiException(error_code=ErrorCodes.NOT_FOUND, message='Data not found')

    header, *data = rows
    total_entries = len(data)

    def process_question(i):
        answers_count = Counter(
            map(lambda row: row[i].strip(), filter(lambda r: len(r) > i and r[i].strip(), data))
        )
        return header[i], answers_count

    question_stats = dict(map(process_question, range(1, len(header))))

    def format_answer(count_list, ingredients):
        fig, ax = plt.subplots(figsize=(6, 3), subplot_kw=dict(aspect="equal"))

        wedges, texts, autotexts = ax.pie(
            count_list, autopct=lambda pct: f"{pct:.1f}%",
            textprops=dict(color="w"), startangle=90)

        ax.legend(wedges, ingredients,
                  title="",
                  loc="center left",
                  bbox_to_anchor=(1, 0, 0.5, 1))

        plt.setp(autotexts, size=8, weight="bold")

        plt.subplots_adjust(left=0.05, right=0.85)

        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', bbox_inches='tight')
        buffer.seek(0)

        image_base64 = base64.b64encode(buffer.read()).decode('utf-8')
        buffer.close()

        return f'<img src="data:image/png;base64,{image_base64}" />'

    def format_question(question_data):
        question, answers = question_data
        answers_count = []
        answers_text = []
        for answer, count in answers.items():
            answers_text.append(f'{answer} ({count} ta)')
            answers_count.append(count)
        answers_html = format_answer(answers_count, answers_text)
        return f"<li><strong>{question}:</strong><br><br><ul>{answers_html}</ul></li>"

    question_html = "".join(map(format_question, question_stats.items()))

    content = format_html(
        """
        <h2>Results / Результаты / Natijalar </h2>
        <p><strong>Total Entries / Общее количество участников / Umumiy qatnashganlar soni:</strong> {}</p>
        <h3>Question-wise Answer Results / Результаты ответов по вопросам / Savollar bo'yicha javoblar natijalari</h3>
        <ul>{}</ul>
        """,
        total_entries,
        format_html(question_html),
    )

    return content
